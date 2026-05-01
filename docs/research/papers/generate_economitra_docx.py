"""Generate Economítra Word document with minimalist cover page."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os

def set_cell_border(cell, **kwargs):
    """Set cell border properties."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        if edge in kwargs:
            element = OxmlElement(f'w:{edge}')
            for k, v in kwargs[edge].items():
                element.set(qn(f'w:{k}'), str(v))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx.enum.text.WD_BREAK.PAGE)

def create_cover_page(doc):
    """Create minimalist cover page."""
    # Add lots of vertical space
    for _ in range(8):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        pf = p.paragraph_format
        pf.line_spacing = Pt(24)

    # Greek title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run('Οικονομίτρα')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.name = 'Georgia'

    # Main title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('Economítra')
    run.font.size = Pt(48)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    run.font.name = 'Georgia'

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('The Measurement of All Things')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = 'Georgia'
    run.font.italic = True

    # Thin line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run('─' * 30)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    # Author
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('Will Glynn')
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    run.font.name = 'Georgia'

    # Year
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('2026')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.name = 'Georgia'

    # Page break
    doc.add_page_break()


def parse_markdown(md_text):
    """Parse markdown into structured sections."""
    lines = md_text.split('\n')
    elements = []
    in_code_block = False
    code_content = []
    i = 0

    while i < len(lines):
        line = lines[i]

        # Skip the front matter (title, subtitle, author) - cover page handles this
        # We start processing after the first '---' separator following the preface marker

        # Code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                elements.append(('code', '\n'.join(code_content)))
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_content.append(line)
            i += 1
            continue

        # Horizontal rules
        if line.strip() == '---':
            i += 1
            continue

        # Headers
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            # Skip the duplicate title/subtitle
            if text in ['Economítra', 'Preface']:
                if text == 'Preface':
                    elements.append(('h1', text))
            else:
                elements.append(('h1', text))
            i += 1
            continue

        if line.startswith('## '):
            text = line[3:].strip()
            if text == 'The Measurement of All Things':
                i += 1
                continue
            elements.append(('h2', text))
            i += 1
            continue

        if line.startswith('### '):
            elements.append(('h3', line[4:].strip()))
            i += 1
            continue

        # Empty lines
        if line.strip() == '':
            i += 1
            continue

        # Regular paragraph (may span multiple lines)
        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() != '' and not lines[i].startswith('#') and not lines[i].startswith('```') and not lines[i].strip() == '---':
            para_lines.append(lines[i])
            i += 1

        text = ' '.join(l.strip() for l in para_lines)
        if text.startswith('**Will Glynn | 2026**'):
            continue
        if text.startswith('*From the Greek'):
            elements.append(('epigraph', text))
            continue
        if text.startswith('*©'):
            elements.append(('footer', text))
            continue
        if text.startswith('*The math doesn'):
            elements.append(('footer', text))
            continue
        if text.startswith('*Economítra is a living'):
            elements.append(('closing', text))
            continue
        if text.startswith('*Economítra.*'):
            elements.append(('closing', text))
            continue

        # Bullet points
        if text.startswith('- '):
            # Re-split since we joined them
            for pl in para_lines:
                stripped = pl.strip()
                if stripped.startswith('- '):
                    elements.append(('bullet', stripped[2:]))
                elif stripped:
                    elements.append(('para', stripped))
            continue

        elements.append(('para', text))

    return elements


def add_formatted_text(paragraph, text, base_size=11, base_color=RGBColor(0x2a, 0x2a, 0x2a)):
    """Add text with inline markdown formatting (bold, italic)."""
    # Process bold+italic, bold, italic, and inline code
    parts = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|`[^`]+`)', text)

    for part in parts:
        if not part:
            continue
        if part.startswith('***') and part.endswith('***'):
            run = paragraph.add_run(part[3:-3])
            run.font.bold = True
            run.font.italic = True
            run.font.size = Pt(base_size)
            run.font.color.rgb = base_color
            run.font.name = 'Georgia'
        elif part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
            run.font.size = Pt(base_size)
            run.font.color.rgb = base_color
            run.font.name = 'Georgia'
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.font.italic = True
            run.font.size = Pt(base_size)
            run.font.color.rgb = base_color
            run.font.name = 'Georgia'
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.size = Pt(base_size - 1)
            run.font.name = 'Consolas'
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        else:
            run = paragraph.add_run(part)
            run.font.size = Pt(base_size)
            run.font.color.rgb = base_color
            run.font.name = 'Georgia'


def build_document(md_path, output_path):
    """Build the Word document."""
    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    # Cover page
    create_cover_page(doc)

    # Epigraph page (the Greek etymology line)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(120)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    text = 'From the Greek economía (household management) and metron (measurement).'
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = 'Georgia'

    doc.add_page_break()

    # Parse and render content
    elements = parse_markdown(md_text)

    section_num = 0
    skip_epigraph = True  # We already rendered the epigraph

    for elem_type, text in elements:
        if elem_type == 'epigraph':
            continue  # Already on its own page

        if elem_type == 'h1':
            # Main section headers (Roman numerals)
            # Clean up: remove leading roman numerals for formatting
            clean = re.sub(r'^[IVXLC]+\.\s*', '', text)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(36)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.keep_with_next = True

            # If it has a roman numeral, show it
            match = re.match(r'^([IVXLC]+)\.\s*(.*)', text)
            if match:
                run = p.add_run(match.group(1) + '\n')
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                run.font.name = 'Georgia'
                run = p.add_run(match.group(2))
                run.font.size = Pt(22)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
                run.font.name = 'Georgia'
            else:
                run = p.add_run(text)
                run.font.size = Pt(22)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
                run.font.name = 'Georgia'

        elif elem_type == 'h2':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = True
            add_formatted_text(p, text, base_size=14, base_color=RGBColor(0x2a, 0x2a, 0x2a))
            for run in p.runs:
                run.font.bold = True

        elif elem_type == 'h3':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            add_formatted_text(p, text, base_size=12, base_color=RGBColor(0x33, 0x33, 0x33))
            for run in p.runs:
                run.font.bold = True
                run.font.italic = True

        elif elem_type == 'code':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.left_indent = Cm(1)
            # Light gray background via shading
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'F5F5F5')
            pPr.append(shd)
            run = p.add_run(text)
            run.font.size = Pt(9)
            run.font.name = 'Consolas'
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        elif elem_type == 'bullet':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(1.2)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            # Add bullet char
            run = p.add_run('  \u2022  ')
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            run.font.name = 'Georgia'
            add_formatted_text(p, text)

        elif elem_type == 'closing':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(8)
            clean = text.strip('*')
            run = p.add_run(clean)
            run.font.size = Pt(12)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            run.font.name = 'Georgia'

        elif elem_type == 'footer':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            clean = text.strip('*')
            run = p.add_run(clean)
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            run.font.name = 'Georgia'

        elif elem_type == 'para':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = Pt(16)
            add_formatted_text(p, text)

    doc.save(output_path)
    print(f"Saved to {output_path}")


if __name__ == '__main__':
    import docx.enum.text
    script_dir = os.path.dirname(os.path.abspath(__file__))
    md_path = os.path.join(script_dir, 'ECONOMITRA.md')
    output_path = os.path.join(script_dir, 'ECONOMITRA.docx')
    build_document(md_path, output_path)
